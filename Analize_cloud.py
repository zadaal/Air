"""
Air data analyze tools
"""
import os.path
import plotly.express as px
px.set_mapbox_access_token(open("mapbox_token.txt").read())
import csv
from datetime import datetime
import datetime as dt
from datetime import date
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import requests
plt.ion()

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

# plt.ioff()

class AnalizeAirData():

    def __init__(self):
        pass

    def calc_period_data(self, df_period, crnt_city, crnt_param):
        timestamp_strt = min(df_period.Timestamp)
        timestamp_end = max(df_period.Timestamp)
        datetime_strt_obj = datetime.fromtimestamp(timestamp_strt)
        datetime_end_obj = datetime.fromtimestamp(timestamp_end)
        datetime_strt = datetime_strt_obj.date().__str__() + '-' + datetime_strt_obj.time().__str__()
        datetime_end = datetime_end_obj.date().__str__() + '-' + datetime_end_obj.time().__str__()
        City = df_period.iloc[0]["City"]
        Id = crnt_city
        param = crnt_param
        mean_val = df_period.value.mean()
        std_val = df_period.value.std()
        median_val = np.median(df_period.value)
        min_val = df_period.value.min()
        max_val = df_period.value.max()
        date_time = df_period["date_time"].iloc[0]
        data_vec = [timestamp_strt, timestamp_end, datetime_strt, datetime_end, City, Id, param, mean_val, std_val,
                    median_val, min_val, max_val, date_time]
        return data_vec


    def plot_stations(self):
        # API_TOKEN=<INSERT_TOKEN_HERE>
        r_by_loc = requests.get('https://air-api.sviva.gov.il/v1/envista/stations',
                                headers={'Authorization': API_TOKEN,
                                         'envi-data-source': 'MANA'}, verify=False)

        json_all_cities = r_by_loc.json()
        df = pd.json_normalize(json_all_cities)
        df['name_reverse'] = df.loc[:, 'name'].apply(lambda x: x[::-1])
        # Load a map:
        fig = px.scatter_mapbox(df, lat="location.latitude", lon="location.longitude", color="active", hover_name="stationId", text="name_reverse",
                                color_continuous_scale=px.colors.cyclical.IceFire, size_max=15, zoom=10)
        fig.show()

        a=1

    def plot_by_city(self, param_name_vec, df_daily, city_codes, y_param, params_by_time, type_str, file_name):
        document = Document()
        document.add_heading(f"{type_str}")
        font = document.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(10)
        section = document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE

        for param in param_name_vec:
            plt.figure()
            ax = plt.gca()
            city_name_lst = list()
            # city_name_lst = list(pd.unique(df["City"]))
            # for city_id in pd.unique(df.Id):
            for city_id, city_name in city_codes.items():
                # city_name=city_codes[city_id]
                if not ((df_daily[(df_daily.param == param) & (df_daily.Id == city_id)]).empty):
                    # df_daily[(df_daily.param==param) & (df_daily.Id==city_id)].plot(x="datetime_strt", y="max", ax=ax)
                    df_daily[(df_daily.param == param) & (df_daily.Id == city_id)].plot(x="date_time", y=y_param, ax=ax, marker='.', lw=0.5)
                    city_name_lst.append(city_name)
                # city_name_lst.append(city_name)

            plt.legend(city_name_lst)
            plt.title(f'{param} by {type_str}, TH: Israel-blue, WHO-red, EU-black')
            if 0:
                line = ax.lines[0]
                time_axis = line.get_xdata()
                israel_th = params_by_time[(params_by_time.param == param)].values[0, 1]
                who_th = params_by_time[(params_by_time.param == param)].values[0, 2]
                eu_th = params_by_time[(params_by_time.param == param)].values[0, 3]
                plt.plot(time_axis, israel_th * np.ones(len(time_axis)), 'b-v')
                plt.plot(time_axis, who_th * np.ones(len(time_axis)), 'r-^')
                plt.plot(time_axis, eu_th * np.ones(len(time_axis)), 'k-o')
            plt.grid()
            plt.show()
            plt.ion()

            for tick in ax.xaxis.get_major_ticks():
                tick.label.set_fontsize(6)
                tick.label.set_rotation('vertical')
            plt.xticks(rotation=45)

            fig_name_to_save = f"{param} by {type_str}.jpeg"
            plt.savefig(fig_name_to_save)

            p = document.add_paragraph(f"{param} \n")
            p.style = document.styles['Normal']
            r1 = p.add_run()
            r1.add_picture(fig_name_to_save, width=Inches(6.8))

            # document.save(f"Air Analize results {type_str}.docx")
            document.save(f"{file_name}.docx")


    def read_csv2dataframe(self, csv_file_name, start_timestamp, end_timestamp):
        data = []

        with open(csv_file_name, newline='') as csvfile:
            csv_data = csv.reader(csvfile, delimiter=',')
            for row in csv_data:
                print(', '.join(row))
                # data.append(', '.join(row))
                data.append(row)

        # Eliminate Empty Values
        data = list(filter(None, data))
        date_timestamp = list()
        new_data = list()
        # Convert date string to time value
        # date_time=list()
        for i_row in range(0, len(data)):
            date_str = data[i_row][0]
            is_data = date_str.find('2023')
            # new_line_data=list()
            if is_data != -1:
                date_time_obj = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%S%z")
                # datetime.datetime.fromtimestamp(0)).total_seconds()
                time_stamp = date_time_obj.timestamp()
                # date_time.append(datetime.fromtimestamp(time_stamp))
                new_line_data = [str(time_stamp)] + data[i_row]
                new_data.append(new_line_data)

        # Convert to Panda's dataframe and convert rellevant fields to float:
        df = pd.DataFrame(new_data, columns=['Timestamp', 'Time', 'City', 'Id', 'param', 'value', 'status', 'valid'])
        df["Timestamp"] = df["Timestamp"].astype(float)
        df["date_time"] = pd.to_datetime(df['Timestamp'] + 2 * 60 * 60, unit='s')
        df["Id"] = df["Id"].astype(float)
        df["value"] = df["value"].astype(float)
        df["status"] = df["status"].astype(float)
        # df["valid"]=df["valid"].astype(boolean)

        # Sort dataframe by time:
        df.sort_values("Timestamp", inplace=True)
        # Remove duplicates:
        df.drop_duplicates(inplace=True)
        # Remove false values:
        df = df[df.valid == 'True']
        min_time_stamp_in_df = df["Timestamp"].min()
        max_time_stamp_in_df = df["Timestamp"].max()

        df = df.reset_index(drop=True)

        start_timestamp = max(min_time_stamp_in_df, start_timestamp)
        end_timestamp = min(max_time_stamp_in_df, end_timestamp)

        param_name_vec = pd.unique(df.param)

        return df, min_time_stamp_in_df, max_time_stamp_in_df, param_name_vec

    def moving_average_threshold(self, df, timeWindow, city, param, threshold, flagPlot):
        # Filter the dataframe to include only the specified city, id, and param
        df_filtered = df[(df['City'] == city) & (df['param'] == param)]

        # Convert the 'Timestamp' column to a datetime object
        # df_filtered['Timestamp'] = pd.to_datetime(df_filtered['date_time'])
        df_filtered.loc[:,'Timestamp'] = pd.to_datetime(df_filtered.loc[:,'date_time'])

        # Sort the dataframe by timestamp
        df_filtered = df_filtered.sort_values('Timestamp')

        # Define a rolling window for the specified time period
        window = pd.Timedelta(seconds=timeWindow)

        # Compute the rolling mean
        rolling_mean = df_filtered.set_index('Timestamp').rolling(window).mean()

        # Filter the rolling mean to include only the timestamps where the mean is above the threshold
        threshold_mask = rolling_mean['value'] > threshold
        rolling_mean_above_threshold = rolling_mean[threshold_mask]

        # Extract the timestamps and values where the rolling mean is above the threshold
        timestamps = rolling_mean_above_threshold.index.strftime('%Y-%m-%d %H:%M:%S')
        values = rolling_mean_above_threshold['value'].tolist()

        if flagPlot:
            plt.figure()
            plt.plot(df_filtered['Timestamp'], df_filtered['value'], 'o-b')
            plt.plot(rolling_mean.index, rolling_mean['value'], 'o-r')
            plt.hlines(threshold, df_filtered['Timestamp'].min(), df_filtered['Timestamp'].max(), label='Threshold',
                       colors='r')
            plt.title(f'{param} in {city} with threshold {threshold}, within {timeWindow/(60*60)} hours')
            plt.grid()
            plt.show()
            plt.ion()


        return timestamps, values

    def run_analyze_T(self, params_air_tbl, city_ids, df, timeWindowStr, city_codes, param_name_vec, T, flagPlot):
        # Get the time window of T seconds
        log_filename = 'log_params.xlsx'
        start_time = datetime.now() - dt.timedelta(seconds=T)
        df_window = df[df['date_time'] >= start_time]
        if timeWindowStr == 'israel_half_hr':
            timeWindowSec = 0.5 * 60 * 60
        elif timeWindowStr == 'israel_hr':
            timeWindowSec = 1 * 60 * 60
        elif timeWindowStr == 'israel_24hr':
            timeWindowSec = 24 * 60 * 60
        elif timeWindowStr == 'israel_8hr':
            timeWindowSec = 8 * 60 * 60
        elif timeWindowStr == 'israel_yr':
            timeWindowSec = 365 * 24 * 60 * 60
        for crnt_param in param_name_vec:
            thresholdValue = params_air_tbl.loc[params_air_tbl['param'] == crnt_param][timeWindowStr].values[0]
            for crnt_city in city_ids:
                # Run the analysis on the time window
                city_name = city_codes.get(crnt_city, 'default_value')
                timestamps, values = self.moving_average_threshold(df_window, timeWindow=timeWindowSec,
                                                                      city=city_name, param=crnt_param,
                                                                      threshold=thresholdValue, flagPlot=flagPlot)
                plt.ion()
                # Save the result to a log file
                if len(values) != 0:
                    log_df = pd.DataFrame({'Report Date': [datetime.now().strftime("%Y-%m-%d")],
                                           'Report Time': [datetime.now().strftime("%H:%M:%S")],
                                           'Time Window': timeWindowStr,
                                           'Station Name': city_name,
                                           'Parameter': crnt_param,
                                           'Time Stamps': [timestamps],
                                           'Values': [values],
                                           'TH Value': [thresholdValue]})
                    if os.path.isfile(log_filename):
                        existing_data = pd.read_excel(log_filename)
                        new_data = pd.concat([existing_data, log_df], ignore_index=True)
                        new_data.to_excel(log_filename, index=False)

                        # # Append data to an existing Excel file with openpyxl
                        # with pd.ExcelWriter(log_filename, engine='openpyxl', mode='a') as writer:
                        #     log_df.to_excel(writer, sheet_name='Sheet2')
                    else:
                        log_df.to_excel(log_filename, index=False)
                        # # Create a new Excel file with xlsxwriter
                        # with pd.ExcelWriter(log_filename, engine='xlsxwriter') as writer:
                        #     log_df.to_excel(writer, sheet_name='Sheet1')

                    # return timestamps, values
