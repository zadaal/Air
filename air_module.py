import requests
import tablib
import time
import pandas as pd
import geopandas as gpd
from shapely.geometry import Point
from shapely import wkt
from urllib3.util.retry import Retry
import plotly.express as px
import folium
import webbrowser
import os
import csv
from datetime import datetime, timezone, date
import datetime as dt
import numpy as np
import matplotlib.pyplot as plt
import requests

from docx import Document # packegae name is python-docx
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
import matplotlib
matplotlib.use("TkAgg")
import io

plt.ion()
"""
Main module with various functions of Air monitoring package
Written By Alon
"""


class data_importer():
    def __init__(self, api_token='ApiToken 1cab20bf-0248-493d-aedc-27aa94445d15'):
        self.api_token = api_token

    def get_stations_info(self):
        try:
            r_by_loc = requests.get('https://air-api.sviva.gov.il/v1/envista/stations',
                                    headers={'Authorization': self.api_token,
                                             'envi-data-source': 'MANA'}, verify=False)
        except requests.exceptions.RequestException as e:
            print(f"An error occurred: {e}")

        json_all_cities = r_by_loc.json()
        json_all_cities_df = pd.DataFrame(json_all_cities)
        json_all_cities_df[["latitude", "longitude"]] = pd.DataFrame(json_all_cities_df["location"].tolist(),
                                                               index=json_all_cities_df.index)
        json_all_cities_df["geometry"] = [
            Point(xy) for xy in zip(json_all_cities_df["longitude"], json_all_cities_df["latitude"])
        ]
        json_all_cities_gdf = gpd.GeoDataFrame(json_all_cities_df, geometry="geometry", crs="EPSG:4326")

        return json_all_cities, json_all_cities_gdf

    def station_data(self, sttn_id):
        r_by_loc = requests.get('https://air-api.sviva.gov.il/v1/envista/stations/{}'.format(str(sttn_id)),
                                headers={'Authorization': self.api_token,
                                         'envi-data-source': 'MANA'}, verify=False)
        json_by_city = r_by_loc.json()
        json_by_city_df = pd.json_normalize(json_by_city)

        return json_by_city, json_by_city_df

    def station_latest_data(self, sttn_id):
        data_processor_instance = data_processor()

        r_by_loc_latest = requests.get('https://air-api.sviva.gov.il/v1/envista/stations/{}/data/latest'.format(str(sttn_id)),
                                headers={'Authorization': self.api_token,
                                         'envi-data-source': 'MANA'}, verify=False)
        json_by_city_latest = r_by_loc_latest.json()
        # json_by_city_latest_df = pd.DataFrame(json_by_city_latest['data'][0])
        json_by_city_latest_df = data_processor_instance.json_to_df(json_by_city_latest)
        # Add city info
        json_by_city, json_by_city_df = self.station_data(sttn_id)
        # Validate location in json_by_city_df
        if (json_by_city_df["location.longitude"][0] == None):
            json_by_city_df[["location.latitude", "location.longitude"]] = (
                json_by_city_df["StationNotes"]
                .str.extract(r"lat\s*:\s*([\d\.]+)\s*long\s*:\s*([\d\.]+)")
                .astype(float)
            )



        joined_df = pd.merge(json_by_city_df, json_by_city_latest_df, how="cross")
        # Convert to geodataframe
        geometry = [Point(xy) for xy in zip(joined_df["location.longitude"], joined_df["location.latitude"])]
        joined_gdf = gpd.GeoDataFrame(joined_df, geometry=geometry, crs="EPSG:4326")

        return json_by_city_latest, joined_gdf

class data_processor():
    def json_to_df(self, data_json):
        data = data_json['data']
        channels_series = pd.DataFrame(data)['channels']
        channels_exploded = channels_series.explode()
        channels_parsed = pd.DataFrame(channels_exploded.tolist())
        # channels_parsed.insert(0, 'stationId', data_json['stationId'])
        # channels_parsed.insert(0, 'date', data[0]['datetime'])

        return channels_parsed


class data_analyzer():

    def __init__(self):
        pass

    def calc_period_data(self, df_period, crnt_city, crnt_param):
        timestamp_strt = min(df_period.Timestamp)
        timestamp_end = max(df_period.Timestamp)
        datetime_strt_obj = datetime.fromtimestamp(timestamp_strt)
        datetime_end_obj = datetime.fromtimestamp(timestamp_end)
        datetime_strt = datetime_strt_obj.date().__str__() + '-' + datetime_strt_obj.time().__str__()
        datetime_end = datetime_end_obj.date().__str__() + '-' + datetime_end_obj.time().__str__()
        City = df_period.iloc[0]["City"]
        Id = crnt_city
        param = crnt_param
        mean_val = df_period.value.mean()
        std_val = df_period.value.std()
        median_val = np.median(df_period.value)
        min_val = df_period.value.min()
        max_val = df_period.value.max()
        date_time = df_period["date_time"].iloc[0]
        data_vec = [timestamp_strt, timestamp_end, datetime_strt, datetime_end, City, Id, param, mean_val, std_val,
                    median_val, min_val, max_val, date_time]
        return data_vec


    def plot_stations(self):
        r_by_loc = requests.get('https://air-api.sviva.gov.il/v1/envista/stations',
                                headers={'Authorization': 'ApiToken 745356f0-5eee-4da8-aa71-b739f4acc081',
                                         'envi-data-source': 'MANA'}, verify=False)

        json_all_cities = r_by_loc.json()
        df = pd.json_normalize(json_all_cities)
        df['name_reverse'] = df.loc[:, 'name'].apply(lambda x: x[::-1])
        # Load a map:
        fig = px.scatter_mapbox(df, lat="location.latitude", lon="location.longitude", color="active", hover_name="stationId", text="name_reverse",
                                color_continuous_scale=px.colors.cyclical.IceFire, size_max=15, zoom=10)
        fig.show()

        a=1

    def plot_by_city(self, param_name_vec, df_daily, city_codes, y_param, params_by_time, type_str, file_name):
        document = Document()
        document.add_heading(f"{type_str}")
        font = document.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(10)
        section = document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE

        for param in param_name_vec:
            plt.figure()
            ax = plt.gca()
            city_name_lst = list()
            # city_name_lst = list(pd.unique(df["City"]))
            # for city_id in pd.unique(df.Id):
            for city_id, city_name in city_codes.items():
                # city_name=city_codes[city_id]
                if not ((df_daily[(df_daily.param == param) & (df_daily.Id == city_id)]).empty):
                    # df_daily[(df_daily.param==param) & (df_daily.Id==city_id)].plot(x="datetime_strt", y="max", ax=ax)
                    df_daily[(df_daily.param == param) & (df_daily.Id == city_id)].plot(x="date_time", y=y_param, ax=ax, marker='.', lw=0.5)
                    city_name_lst.append(city_name)
                # city_name_lst.append(city_name)

            plt.legend(city_name_lst)
            plt.title(f'{param} by {type_str}, TH: Israel-blue, WHO-red, EU-black')
            if 0:
                line = ax.lines[0]
                time_axis = line.get_xdata()
                israel_th = params_by_time[(params_by_time.param == param)].values[0, 1]
                who_th = params_by_time[(params_by_time.param == param)].values[0, 2]
                eu_th = params_by_time[(params_by_time.param == param)].values[0, 3]
                plt.plot(time_axis, israel_th * np.ones(len(time_axis)), 'b-v')
                plt.plot(time_axis, who_th * np.ones(len(time_axis)), 'r-^')
                plt.plot(time_axis, eu_th * np.ones(len(time_axis)), 'k-o')
            plt.grid()
            plt.show()
            plt.ion()

            for tick in ax.xaxis.get_major_ticks():
                tick.label.set_fontsize(6)
                tick.label.set_rotation('vertical')
            plt.xticks(rotation=45)

            fig_name_to_save = f"{param} by {type_str}.jpeg"
            plt.savefig(fig_name_to_save)

            p = document.add_paragraph(f"{param} \n")
            p.style = document.styles['Normal']
            r1 = p.add_run()
            r1.add_picture(fig_name_to_save, width=Inches(6.8))

            # document.save(f"Air Analize results {type_str}.docx")
            document.save(f"{file_name}.docx")


    def read_csv2dataframe(self, csv_file_name, start_timestamp, end_timestamp):
        data = []

        with open(csv_file_name, newline='') as csvfile:
            csv_data = csv.reader(csvfile, delimiter=',')
            for row in csv_data:
                print(', '.join(row))
                # data.append(', '.join(row))
                data.append(row)

        # Eliminate Empty Values
        data = list(filter(None, data))
        date_timestamp = list()
        new_data = list()
        # Convert date string to time value
        # date_time=list()
        for i_row in range(0, len(data)):
            date_str = data[i_row][0]
            is_data = date_str.find('2023')
            # new_line_data=list()
            if is_data != -1:
                date_time_obj = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%S%z")
                # datetime.datetime.fromtimestamp(0)).total_seconds()
                time_stamp = date_time_obj.timestamp()
                # date_time.append(datetime.fromtimestamp(time_stamp))
                new_line_data = [str(time_stamp)] + data[i_row]
                new_data.append(new_line_data)

        # Convert to Panda's dataframe and convert rellevant fields to float:
        df = pd.DataFrame(new_data, columns=['Timestamp', 'Time', 'City', 'Id', 'param', 'value', 'status', 'valid'])
        df["Timestamp"] = df["Timestamp"].astype(float)
        df["date_time"] = pd.to_datetime(df['Timestamp'] + 2 * 60 * 60, unit='s')
        df["Id"] = df["Id"].astype(float)
        df["value"] = df["value"].astype(float)
        df["status"] = df["status"].astype(float)
        # df["valid"]=df["valid"].astype(boolean)

        # Sort dataframe by time:
        df.sort_values("Timestamp", inplace=True)
        # Remove duplicates:
        df.drop_duplicates(inplace=True)
        # Remove false values:
        df = df[df.valid == 'True']
        min_time_stamp_in_df = df["Timestamp"].min()
        max_time_stamp_in_df = df["Timestamp"].max()

        df = df.reset_index(drop=True)

        start_timestamp = max(min_time_stamp_in_df, start_timestamp)
        end_timestamp = min(max_time_stamp_in_df, end_timestamp)

        param_name_vec = pd.unique(df.param)

        return df, min_time_stamp_in_df, max_time_stamp_in_df, param_name_vec

    def read_csv2df(self, csv_path, start_datetime=None, end_datetime=None, export_to_excel=False):
        # Read the entire file.
        with open(csv_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # Filter out lines that repeat the header "Time,City,Id,parameter,value,status,valid".
        filtered_lines = [
            line for line in lines
            if not line.strip().startswith("Time,City,Id,parameter,value,status,valid")
               and line.strip() != ""
        ]

        # Convert the filtered lines back to a single string so we can pass it to read_csv.
        filtered_data_str = "\n".join(filtered_lines)

        # Read the CSV data using pandas. Since we've removed the header lines,
        # we can supply our own column names:
        df = pd.read_csv(
            io.StringIO(filtered_data_str),
            header=None,  # tells pandas that there is no header line in the data
            names=["Time", "City", "Id", "Parameter", "Value", "Status", "Valid"]
        )

        # Convert the "Time" column to a datetime (pandas will handle the ISO-like format).
        df["Time"] = pd.to_datetime(df["Time"])
        df['Time'] = df['Time'].dt.tz_localize(None)  # convert it to UTC
        df["date_time"] = df["Time"].apply(lambda x: x.timestamp())
        if (start_datetime is not None) and (end_datetime is not None):
            df = df[(df['Time'] > start_datetime) & (df['Time'] < end_datetime)]

        # Remove unvalid data:
        df = df[df['Valid'] == True]

        # Import stations info and join dfs
        data_importer_inst = data_importer()
        all_cities_json, all_cities_df = data_importer_inst.get_stations_info()
        all_cities_df['Id'] = all_cities_df['stationId']
        df_merged = pd.merge(df, all_cities_df, on='Id', how='left')

        # Convert to geodataframe:
        gdf_merged = gpd.GeoDataFrame(df_merged, geometry='geometry')
        gdf_merged.set_crs(epsg=4326, inplace=True)
        # Export df to excel:
        if export_to_excel:
            output_file = 'window_data.xlsx'
            gdf_merged.to_excel(output_file, index=False)
        return gdf_merged

    def moving_average_threshold(self, df, timeWindow, city, param, threshold, flagPlot):
        """
        df : DataFrame that has columns ['Time', 'City', 'Parameter', 'Value']
             and where 'Time' is already a datetime type.
        timeWindow : window size in seconds for the rolling mean
        city : city name to filter
        param : parameter name to filter
        threshold : threshold to compare the rolling mean against
        flagPlot : boolean indicating whether to show a matplotlib plot
        """
        # 1. Filter the DataFrame
        df_filtered = df[(df['City'] == city) & (df['Parameter'] == param)].copy()
        if df_filtered.empty:
            return None, []

        # 2. Sort by "Time"
        df_filtered = df_filtered.sort_values('Time')

        # 3. Set "Time" as the index to enable time-based rolling
        df_filtered = df_filtered.set_index('Time')

        # 4. Convert the timeWindow (in seconds) to a pandas Timedelta
        window = pd.Timedelta(seconds=timeWindow)

        # 5. Compute the rolling mean for the "Value" column
        rolling_mean = df_filtered['Value'].rolling(window=window).mean()

        # 6. Identify rows where the rolling mean is above the threshold
        threshold_mask = rolling_mean > threshold
        rolling_mean_above_threshold = rolling_mean[threshold_mask]

        # 7. Extract timestamps and values where above threshold
        timestamps = rolling_mean_above_threshold.index.strftime('%Y-%m-%d %H:%M:%S')
        values = rolling_mean_above_threshold.tolist()

        # Optional: plotting
        if flagPlot:
            plt.figure()
            plt.plot(df_filtered.index, df_filtered['Value'], 'o-b', label='Original Values')
            plt.plot(rolling_mean.index, rolling_mean, 'o-r', label='Rolling Mean')
            plt.axhline(threshold, linestyle='--', label='Threshold', color='r')
            plt.title(f'{param} in {city}, threshold {threshold}, window={timeWindow / 3600} hours')
            plt.legend()
            plt.grid()
            plt.show(block=True)
            # plt.show()
            # plt.ion()

        return timestamps, values

    def run_analyze_T(self, params_air_tbl, city_ids, df, timeWindowStr, city_codes, param_name_vec, T, flagPlot):
        plotter = data_plotter()

        # Get the time window of T seconds
        log_filename = 'log_params.xlsx'
        start_time = datetime.now() - dt.timedelta(seconds=T)
        df_window = df[df['Time'] >= start_time.strftime("%Y-%m-%dT%H:%M:%S%z")]
        if timeWindowStr == 'israel_half_hr':
            timeWindowSec = 0.5 * 60 * 60
        elif timeWindowStr == 'israel_hr':
            timeWindowSec = 1 * 60 * 60
        elif timeWindowStr == 'israel_24hr':
            timeWindowSec = 24 * 60 * 60
        elif timeWindowStr == 'israel_8hr':
            timeWindowSec = 8 * 60 * 60
        elif timeWindowStr == 'israel_yr':
            timeWindowSec = 365 * 24 * 60 * 60
        for crnt_param in param_name_vec:
            thresholdValue = params_air_tbl.loc[params_air_tbl['param'] == crnt_param][timeWindowStr].values[0]
            for crnt_city in city_ids:
                # Run the analysis on the time window
                city_name = city_codes.get(crnt_city, 'default_value')
                if crnt_param=='WD':
                    wd = df_window[(df_window['Parameter'] == crnt_param) & (df_window['City'] == city_name)]['Value'].to_numpy()
                    tw = df_window[(df_window['Parameter'] == crnt_param) & (df_window['City'] == city_name)]['Time']
                    if len(wd) != 0:
                        min_timestamp_str = tw.min().strftime('%Y-%m-%d %H:%M:%S')
                        max_timestamp_str = tw.max().strftime('%Y-%m-%d %H:%M:%S')
                        plotter.plot_polar_wind_hist(wind_direction_deg=wd, city_name=city_name, start_time=min_timestamp_str, end_time = max_timestamp_str)
                else:
                    timestamps, values = self.moving_average_threshold(df_window, timeWindow=timeWindowSec,
                                                                       city=city_name, param=crnt_param,
                                                                       threshold=thresholdValue, flagPlot=flagPlot)
                    plt.ion()
                    # Save the result to a log file
                    if len(values) != 0:
                        log_df = pd.DataFrame({'Report Date': [datetime.now().strftime("%Y-%m-%d")],
                                               'Report Time': [datetime.now().strftime("%H:%M:%S")],
                                               'Time Window': timeWindowStr,
                                               'Station Name': city_name,
                                               'Parameter': crnt_param,
                                               'Time Stamps': [timestamps],
                                               'Values': [values],
                                               'TH Value': [thresholdValue]})
                        if os.path.isfile(log_filename):
                            existing_data = pd.read_excel(log_filename)
                            new_data = pd.concat([existing_data, log_df], ignore_index=True)
                            new_data.to_excel(log_filename, index=False)

                            # # Append data to an existing Excel file with openpyxl
                            # with pd.ExcelWriter(log_filename, engine='openpyxl', mode='a') as writer:
                            #     log_df.to_excel(writer, sheet_name='Sheet2')
                        else:
                            log_df.to_excel(log_filename, index=False)
                            # # Create a new Excel file with xlsxwriter
                            # with pd.ExcelWriter(log_filename, engine='xlsxwriter') as writer:
                            #     log_df.to_excel(writer, sheet_name='Sheet1')

                        # return timestamps, values

class data_plotter():
    def plot_stations(self, stations_gdf):
        fig = px.scatter_mapbox(
            stations_gdf,
            lat="latitude",
            lon="longitude",
            hover_name='stationId',  # optional: which column to display in hover
            hover_data = {
                'name': True,
                'active': True,
                'owner': True,
            },
            # color="red",  # optional: color points by a column
            zoom=10,
            mapbox_style="carto-positron"  #"open-street-map" or "carto-positron", "stamen-terrain", etc.
        )

        fig.show()

    def plot_stations_folium(self, stations_gdf):
        stations_gdf = stations_gdf.drop(columns=["monitors"])
        map_widget = stations_gdf.explore(tiles="CartoDB positron")
        map_widget.save("stations.html")
        webbrowser.open("stations.html")

    def plot_polar_wind_hist(self, wind_direction_deg, city_name, start_time=None, end_time=None):
        if len(wind_direction_deg) != 0:
            # Convert direction so the arrows shows the destination rather than source:
            wind_direction_corrected_deg = (wind_direction_deg + 180) % 360  # Ensure 0-360 degrees
            # Convert to radians for plotting
            wind_direction_corrected_rad = np.deg2rad(wind_direction_corrected_deg)
            num_bins = 36  # 10-degree bins
            bins = np.linspace(0, 2 * np.pi, num_bins + 1)  # Bins in radians

            # Compute histogram (count occurrences in each bin)
            hist, _ = np.histogram(wind_direction_corrected_rad, bins=bins)

            # Normalize (optional)
            hist = hist / np.max(hist)  # Normalize to max value

            # Plot
            fig, ax = plt.subplots(subplot_kw={'projection': 'polar'})
            ax.set_theta_zero_location('N')  # North is at the top
            ax.set_theta_direction(-1)  # Clockwise angles

            # Plot histogram bars
            bars = ax.bar(bins[:-1], hist, width=(2 * np.pi / num_bins), bottom=0.0, color='b', alpha=0.6,
                          edgecolor='black')

            # Labels
            if start_time==None:
                ax.set_title("Winds directions at " + city_name, fontsize=14)
            else:
                ax.set_title("Winds directions at " + city_name + ' from: ' + start_time + ' to: ' + end_time, fontsize=14)
            plt.show()

