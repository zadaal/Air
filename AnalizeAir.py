import csv
from datetime import datetime
from datetime import date
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
plt.ion()

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.section import WD_ORIENT


# plt.ioff()
# Use station IDs:
# 7 - Beit Shemesh (line #5)
# 32 - rehovot (line #12)
# 64 - Modiin (line #26)
# 76 - Carmei Yosef (line #34)
# 78 Achisamach (line #35)
# 367 Beit Hashmonai (line #111)
# 397 Mobile (line #122)
# 338 Yad Rambam

city_codes = {
    31: 'Modiin Hinanit',
    40: 'Yad Rambam 1',
    64: 'Modiin',
    76: 'Carmei Yosef',
    77: 'Kfar Shmuel',
    78: 'Achisamach',
    193: 'Omanim Ramla',
    367: 'Beit Hashmonai',
    338: 'Yad Rambam New'
}

def CalcPeriodData(df_period, crnt_city, crnt_param):
    timestamp_strt = min(df_period.Timestamp)
    timestamp_end = max(df_period.Timestamp)
    datetime_strt_obj=datetime.fromtimestamp(timestamp_strt)
    datetime_end_obj=datetime.fromtimestamp(timestamp_end)
    datetime_strt = datetime_strt_obj.date().__str__() + '-' + datetime_strt_obj.time().__str__()
    datetime_end = datetime_end_obj.date().__str__() + '-' + datetime_end_obj.time().__str__()
    City = df_period.iloc[0]["City"]
    Id = crnt_city
    param = crnt_param
    mean_val = df_period.value.mean()
    std_val = df_period.value.std()
    median_val = np.median(df_period.value)
    min_val = df_period.value.min()
    max_val = df_period.value.max()
    date_time = df_period["date_time"].iloc[0]
    data_vec = [timestamp_strt, timestamp_end,datetime_strt,datetime_end, City, Id, param, mean_val, std_val, median_val, min_val, max_val,date_time]
    return data_vec

def PlotByCity(param_name_vec, df_daily, city_codes, y_param, params_by_time, type_str):
    document = Document()
    document.add_heading(f"Air Analize results {type_str}")
    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(10)
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE

    for param in param_name_vec:
        plt.figure()
        ax = plt.gca()
        city_name_lst=list()
        # city_name_lst = list(pd.unique(df["City"]))
        # for city_id in pd.unique(df.Id):
        for city_id, city_name in city_codes.items():
            # city_name=city_codes[city_id]
            if not ((df_daily[(df_daily.param == param) & (df_daily.Id == city_id)]).empty):
                # df_daily[(df_daily.param==param) & (df_daily.Id==city_id)].plot(x="datetime_strt", y="max", ax=ax)
                df_daily[(df_daily.param == param) & (df_daily.Id == city_id)].plot(x="date_time", y=y_param, ax=ax,marker='.')
                city_name_lst.append(city_name)
        line = ax.lines[0]
        time_axis = line.get_xdata()
        plt.legend(city_name_lst)
        plt.title(f'{param} by {type_str}, TH: Israel-blue, WHO-red, EU-black')
        israel_th = params_by_time[(params_by_time.param == param)].values[0, 1]
        who_th = params_by_time[(params_by_time.param == param)].values[0, 2]
        eu_th = params_by_time[(params_by_time.param == param)].values[0, 3]
        plt.plot(time_axis, israel_th*np.ones(len(time_axis)),'b-v')
        plt.plot(time_axis, who_th * np.ones(len(time_axis)), 'r-^')
        plt.plot(time_axis, eu_th * np.ones(len(time_axis)), 'k-o')
        plt.grid()
        plt.show(block=False)


        for tick in ax.xaxis.get_major_ticks():
            tick.label.set_fontsize(6)
            tick.label.set_rotation('vertical')
        plt.xticks(rotation=45)

        fig_name_to_save=f"{param} by {type_str}.jpeg"
        plt.savefig(fig_name_to_save)

        p = document.add_paragraph(f"{param} \n")
        p.style = document.styles['Normal']
        r1 = p.add_run()
        r1.add_picture(fig_name_to_save, width=Inches(6.8))

    document.save(f"Air Analize results {type_str}.docx")

params_air=pd.read_excel('params_air.xlsx')

start_date_time=datetime(2025,1,1,0,0,0)
end_date_time=datetime(2025,12,31,23,30,0)

start_timestamp=datetime.timestamp(start_date_time)
end_timestamp=datetime.timestamp(end_date_time)


data=[]

with open('output.csv', newline='') as csvfile:
    csv_data = csv.reader(csvfile, delimiter=',')
    for row in csv_data:
        print(', '.join(row))
        # data.append(', '.join(row))
        data.append(row)

# Eliminate Empty Values
data=list(filter(None,data))
date_timestamp=list()
new_data=list()
# Convert date string to time value
# date_time=list()
for i_row in range(0,len(data)):
    date_str=data[i_row][0]
    is_data=date_str.find('2021')
    # new_line_data=list()
    if is_data!=-1:
        date_time_obj = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%S%z")
        # datetime.datetime.fromtimestamp(0)).total_seconds()
        time_stamp=date_time_obj.timestamp()
        # date_time.append(datetime.fromtimestamp(time_stamp))
        new_line_data=[str(time_stamp)] + data[i_row]
        new_data.append(new_line_data)

# Convert to Panda's dataframe and convert rellevant fields to float:
df=pd.DataFrame(new_data,columns=['Timestamp','Time','City','Id','param','value','status','valid'])
df["Timestamp"]=df["Timestamp"].astype(float)
df["date_time"] = pd.to_datetime(df['Timestamp']+2*60*60, unit='s')
df["Id"]=df["Id"].astype(float)
df["value"]=df["value"].astype(float)
df["status"]=df["status"].astype(float)
# df["valid"]=df["valid"].astype(boolean)

# Sort dataframe by time:
df.sort_values("Timestamp", inplace=True)
# Remove duplicates:
df.drop_duplicates(inplace=True)
# Remove false values:
df = df[df.valid == 'True']
min_time_stamp_in_df = df["Timestamp"].min()
max_time_stamp_in_df = df["Timestamp"].max()

start_timestamp = max(min_time_stamp_in_df, start_timestamp)
end_timestamp = min(max_time_stamp_in_df, end_timestamp)

param_name_vec=pd.unique(df.param)
t_win_vec_daily=np.arange(start_timestamp,end_timestamp+24*60*60,24*60*60)
t_win_vec_weekly=np.arange(start_timestamp,end_timestamp,7*24*60*60)

df_daily=pd.DataFrame([],columns= ('timestamp_strt','timestamp_end','datetime_strt','datetime_end','City','Id','param','mean_val','std_val','median_val','min','max','date_time'))
df_weekly=pd.DataFrame([],columns= ('timestamp_strt','timestamp_end','datetime_strt','datetime_end','City','Id','param','mean_val','std_val','median_val','min','max','date_time'))

for crnt_param in param_name_vec:
    for crnt_city in pd.unique(df.Id):
        for i_t_daily in range(0,len(t_win_vec_daily)-1):
            crnt_daily_data=df[(df.Timestamp>=t_win_vec_daily[i_t_daily]) & (df.Timestamp<t_win_vec_daily[i_t_daily+1]) & (df.param==crnt_param) & (df.Id==crnt_city) & (df.valid=='True')]
            if not(crnt_daily_data.empty):
                data_vec=CalcPeriodData(crnt_daily_data, crnt_city, crnt_param)

                crnt_df_daily=pd.DataFrame([data_vec],columns= ('timestamp_strt','timestamp_end','datetime_strt','datetime_end','City','Id','param','mean_val','std_val','median_val','min','max','date_time'))
                df_daily=df_daily.append(crnt_df_daily)
                b=1
        for i_t_weekly in range(0,len(t_win_vec_weekly)-1):
            crnt_weekly_data = df[
                (df.Timestamp > t_win_vec_daily[i_t_weekly]) & (df.Timestamp <= t_win_vec_daily[i_t_weekly + 1]) & (
                            df.param == crnt_param) & (df.Id == crnt_city) & (df.valid == 'True')]
            if not (crnt_weekly_data.empty):
                data_vec = CalcPeriodData(crnt_weekly_data, crnt_city, crnt_param)
                crnt_df_weekly = pd.DataFrame([data_vec], columns=(
                'timestamp_strt', 'timestamp_end', 'datetime_strt','datetime_end','City', 'Id', 'param', 'mean_val', 'std_val', 'median_val', 'min', 'max','date_time'))
                df_weekly = df_weekly.append(crnt_df_weekly)

                c=1


# Plot results in dataframes:
cond = {'param': 'NOX', 'Id': 367}

# PlotByCity(param_name_vec, df_daily, city_codes, 'max', params_air)

params_by_time_half_hr = params_air[["param","israel_half_hr","who_half_hr","eu_half_hr"]]
PlotByCity(param_name_vec, df, city_codes, 'value', params_by_time_half_hr,'Half Hour')

# params_by_time_daily = params_air[["param","israel_24hr","who_24hr","eu_24hr"]]
# PlotByCity(param_name_vec, df_daily, city_codes, 'max', params_by_time_daily,'24hr max')

a=1

